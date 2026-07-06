"""
DOCX 文件处理模块
将 DOCX 转换为 Markdown

优先使用 python-docx 直接提取文本并识别标题结构，
保留 mammoth + markdownify 作为备选方案（处理含复杂表格/图片的文件）。
"""

import asyncio
import re
from pathlib import Path
from typing import Optional


# 标题识别正则模式
# 匹配 "第X条"、"第一章" 等中文序号标题
_CHAPTER_PATTERN = re.compile(
    r'^(第[一二三四五六七八九十百千零\d]+[章节条款部分篇部]|[一二三四五六七八九十]+[、.])'
)
# 匹配 "1."、"1.1"、"1.1.1" 等数字编号标题
_NUMERIC_PATTERN = re.compile(
    r'^(\d{1,2}[\.、]\s|\d{1,2}\.\d{1,2}[\s\.、]|\d{1,2}\.\d{1,2}\.\d{1,2}[\s\.、])'
)
# 匹配 "(一)"、"（二）" 等括号序号
_BRACKET_PATTERN = re.compile(
    r'^[（(][一二三四五六七八九十]+[）)]'
)


def _detect_heading_level(text: str) -> int:
    """
    根据文本模式推断标题层级。

    Returns:
        0 = 非标题（普通段落）
        1 = 一级标题（第X章/条、1.）
        2 = 二级标题（1.1）
        3 = 三级标题（1.1.1、(一)）
    """
    stripped = text.strip()
    if not stripped:
        return 0

    # 短行 + 中文章节号 → 一级
    if _CHAPTER_PATTERN.match(stripped):
        # "第X条" 通常是二级，"第X章" 才是一级
        if re.match(r'^第[一二三四五六七八九十百千零\d]+章', stripped):
            return 1
        if re.match(r'^第[一二三四五六七八九十百千零\d]+条', stripped):
            return 2
        return 1

    # 三级编号 "1.1.1" → 三级标题
    if re.match(r'^\d{1,2}\.\d{1,2}\.\d{1,2}[\s\.、]', stripped):
        return 4

    # 二级编号 "1.1" → 三级标题（在"第X条"之下）
    if re.match(r'^\d{1,2}\.\d{1,2}[\s\.、]', stripped):
        return 3

    # 一级编号 "1." → 二级标题（仅在行较短时）
    if re.match(r'^\d{1,2}[\.、]\s', stripped):
        # 如果是 "1.1" 会被上面匹配，这里只匹配 "1." 或 "1、"
        # 且行不能太长（标题通常较短）
        if len(stripped) < 80:
            return 3
        return 0

    # 括号序号 "(一)" → 三级标题
    if _BRACKET_PATTERN.match(stripped):
        return 3

    return 0


def _is_heading_style(style_name: str) -> int:
    """
    根据 Word 段落样式判断标题层级。
    Heading 1 → 1, Heading 2 → 2, ...
    """
    if not style_name:
        return 0
    style_lower = style_name.lower()
    if 'heading' in style_lower or '标题' in style_name:
        # 提取数字
        match = re.search(r'(\d+)', style_name)
        if match:
            return min(int(match.group(1)), 6)
        return 1
    # TOC 样式不算标题
    return 0


def docx_to_markdown(file_path: str) -> str:
    """
    使用 python-docx 将 DOCX 转为结构化 Markdown。

    - 无表格时：python-docx 直接提取，识别标题层级
    - 有表格时：用 mammoth 转 HTML（表格完整），再对 HTML 中的标题行补 # 标记

    Args:
        file_path: DOCX 文件路径

    Returns:
        Markdown 文本
    """
    from docx import Document

    doc = Document(file_path)

    # 如果文档包含表格，使用 mammoth 处理（表格更完整），
    # 然后对 mammoth 输出的纯文本补标题标记
    if doc.tables:
        return _docx_with_tables(file_path, doc)

    # 无表格：直接用 python-docx 提取，保留标题结构
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        # 优先用 Word 样式判断标题
        style_level = _is_heading_style(para.style.name)

        # DOCX 可能把多行内容放在一个段落里（用 \n 分隔）
        # 需要先按行拆分，逐行识别标题
        sub_lines = text.split('\n')

        for sub_text in sub_lines:
            sub_text = sub_text.strip()
            if not sub_text:
                lines.append('')
                continue

            if style_level > 0:
                # 有明确的标题样式
                lines.append(f'{"#" * style_level} {sub_text}')
            else:
                # 样式为 Normal，用文本模式推断
                detected_level = _detect_heading_level(sub_text)
                if detected_level > 0:
                    lines.append(f'{"#" * detected_level} {sub_text}')
                else:
                    lines.append(sub_text)

    # 清理连续空行（最多保留两个换行）
    result = '\n'.join(lines)
    result = re.sub(r'\n{3,}', '\n\n', result)

    return result.strip() + '\n'


def _docx_with_tables(file_path: str, doc) -> str:
    """
    处理含表格的 DOCX：用 mammoth 转 HTML（表格完整），
    再用 markdownify 转 Markdown，最后对标题行补 # 标记。

    Args:
        file_path: DOCX 文件路径
        doc: python-docx Document 对象（用于提取标题信息）
    """
    # 1. mammoth → HTML → Markdown（表格完整）
    html_content = docx_to_html(file_path)
    md_content = html_to_markdown(html_content)

    # 2. 对 Markdown 中的标题行补 # 标记
    #    mammoth 输出的标题可能没有 <h> 标签，导致 markdownify 输出纯文本
    #    按行扫描，用文本模式推断标题层级并补标记
    lines = md_content.split('\n')
    result_lines = []
    for line in lines:
        stripped = line.strip()
        # 已经有 # 标记的跳过
        if stripped.startswith('#'):
            result_lines.append(line)
            continue
        # 推断标题层级
        detected = _detect_heading_level(stripped)
        if detected > 0:
            result_lines.append(f'{"#" * detected} {stripped}')
        else:
            result_lines.append(line)

    result = '\n'.join(result_lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip() + '\n'


async def process_docx(file_path: str, output_path: str, **kwargs) -> dict:
    """
    将 DOCX 文件转换为 Markdown

    流程:
    1. 优先使用 python-docx 直接提取结构化 Markdown（识别标题层级）
    2. 备选：使用 mammoth + markdownify（处理复杂格式）

    Args:
        file_path: DOCX 文件路径
        output_path: 输出 Markdown 文件路径
        **kwargs: 额外参数

    Returns:
        dict: 处理结果 {'success': bool, 'output_path': str}
    """
    try:
        print(f'开始转换 DOCX: {file_path}')

        loop = asyncio.get_event_loop()

        # 1. 优先使用 python-docx 直接转 Markdown（保留标题结构）
        try:
            markdown_content = await loop.run_in_executor(
                None,
                lambda: docx_to_markdown(file_path)
            )
        except Exception as e:
            print(f'python-docx 转换失败，尝试 mammoth: {e}')
            # 2. 备选方案：mammoth + markdownify
            markdown_content = await loop.run_in_executor(
                None,
                lambda: _mammoth_fallback(file_path)
            )

        # 3. 写入输出文件
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(markdown_content, encoding='utf-8')

        print(f'DOCX 转换完成：{output_path}')

        return {
            'success': True,
            'output_path': str(output_path),
            'format': 'docx'
        }

    except Exception as e:
        print(f'DOCX 转换失败：{e}')
        return {
            'success': False,
            'error': str(e),
            'format': 'docx'
        }


def _mammoth_fallback(file_path: str) -> str:
    """mammoth + markdownify 备选方案"""
    html_content = docx_to_html(file_path)
    return html_to_markdown(html_content)


def docx_to_html(file_path: str) -> str:
    """
    使用 mammoth 将 DOCX 转为 HTML
    """
    try:
        import mammoth

        with open(file_path, 'rb') as f:
            result = mammoth.convert_to_html(f)

        if result.messages:
            print(f'mammoth 警告：{result.messages}')

        return result.value

    except ImportError:
        raise ImportError("需要安装 mammoth: pip install mammoth")


def html_to_markdown(html_content: str) -> str:
    """
    使用 markdownify 将 HTML 转为 Markdown
    """
    try:
        from markdownify import markdownify as md

        return md(
            html_content,
            heading_style='ATX',
            bullets='-',
            escape_asterisks=False,
            escape_backslashes=False,
            table_infer_header=True
        )

    except ImportError:
        try:
            import html2text

            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            h.ignore_tables = False
            h.table_ignore = 0
            h.mark_code = False
            return h.handle(html_content)

        except ImportError:
            raise ImportError(
                "需要安装 markdownify 或 html2text:\n"
                "  pip install markdownify\n"
                "  或\n"
                "  pip install html2text"
            )


async def process_docx_advanced(
    file_path: str,
    output_path: str,
    preserve_styles: bool = False,
    **kwargs
) -> dict:
    """
    高级 DOCX 处理 (可选保留样式信息)
    """
    try:
        markdown_content = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: docx_to_markdown(file_path)
        )

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(markdown_content, encoding='utf-8')

        return {
            'success': True,
            'output_path': str(output_path),
            'format': 'docx',
            'preserve_styles': preserve_styles
        }

    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'format': 'docx'
        }
